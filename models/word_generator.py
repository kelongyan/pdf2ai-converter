"""Word 文档生成器 - 基于 AI 视觉模型"""

from typing import List, Dict, Any
from pathlib import Path
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from tqdm import tqdm
from .base import VisionModel


class WordGenerator:
    """基于 AI 视觉模型的 Word 文档生成器

    使用视觉大模型分析 PDF 页面的布局和样式，
    生成尽可能保持原样排版的 Word 文档。
    """

    def __init__(self, vision_model: VisionModel, mode: str = "precise"):
        """初始化生成器

        Args:
            vision_model: 视觉模型实例
            mode: 转换模式
                - "fast": 快速模式，只保留内容结构
                - "precise": 精确模式，尽可能保留样式
        """
        self.vision_model = vision_model
        self.mode = mode
        self.logger = None  # 可选：添加日志

    def process_page_to_word(self, image_path: str, page_num: int) -> Dict[str, Any]:
        """处理单页 PDF，返回结构化数据

        Args:
            image_path: 页面图片路径
            page_num: 页码

        Returns:
            包含页面元素和样式的字典
        """
        if self.mode == "fast":
            return self._process_page_fast(image_path, page_num)
        else:
            return self._process_page_precise(image_path, page_num)

    def _process_page_fast(self, image_path: str, page_num: int) -> Dict[str, Any]:
        """快速模式：只提取内容结构"""
        prompt = """分析这页 PDF 的内容，返回 JSON 格式：

{
  "page": 1,
  "elements": [
    {
      "type": "heading",
      "level": 1,
      "content": "标题文本"
    },
    {
      "type": "paragraph",
      "content": "段落文本"
    },
    {
      "type": "table",
      "rows": [
        ["单元格1", "单元格2"],
        ["单元格3", "单元格4"]
      ]
    },
    {
      "type": "list",
      "items": ["项目1", "项目2"]
    }
  ]
}

要求：
1. 按从上到下的顺序识别所有元素
2. 准确区分标题层级（1-6）
3. 表格要包含完整的行列数据
4. 列表要区分有序/无序
5. 只返回 JSON，不要其他说明

直接输出 JSON："""

        try:
            response = self.vision_model.process_page(image_path, prompt)
            # 清理可能的 markdown 代码块标记
            response = response.strip()
            if response.startswith("```"):
                response = re.sub(r'^```json?\s*\n', '', response)
                response = re.sub(r'\n```\s*$', '', response)

            data = json.loads(response)
            data["page"] = page_num
            return data
        except json.JSONDecodeError as e:
            print(f"⚠️  第 {page_num} 页 JSON 解析失败: {e}")
            print(f"原始响应: {response[:200]}...")
            return {
                "page": page_num,
                "elements": [
                    {"type": "paragraph", "content": f"[第 {page_num} 页处理失败]"}
                ]
            }
        except Exception as e:
            print(f"❌ 第 {page_num} 页处理失败: {e}")
            return {
                "page": page_num,
                "elements": [
                    {"type": "paragraph", "content": f"[第 {page_num} 页处理失败: {e}]"}
                ]
            }

    def _process_page_precise(self, image_path: str, page_num: int) -> Dict[str, Any]:
        """精确模式：提取内容和样式"""
        prompt = """分析这页 PDF 的布局、内容和样式，返回 JSON 格式：

{
  "page": 1,
  "layout": "single_column",
  "elements": [
    {
      "type": "heading",
      "level": 1,
      "content": "标题文本",
      "style": {
        "font": "宋体",
        "size": 16,
        "bold": true,
        "color": "#000000",
        "align": "center"
      }
    },
    {
      "type": "paragraph",
      "content": "段落文本",
      "style": {
        "font": "宋体",
        "size": 12,
        "bold": false,
        "italic": false,
        "color": "#000000",
        "align": "justify"
      }
    },
    {
      "type": "table",
      "rows": [
        ["表头1", "表头2"],
        ["数据1", "数据2"]
      ],
      "style": {
        "has_header": true,
        "border": true
      }
    }
  ]
}

要求：
1. layout: "single_column" 或 "two_column"
2. 准确识别字体名称（宋体、黑体、Times New Roman 等）
3. 字号单位为磅（pt）
4. 颜色用十六进制表示（#RRGGBB）
5. align: "left", "center", "right", "justify"
6. 只返回 JSON，不要其他说明

直接输出 JSON："""

        try:
            response = self.vision_model.process_page(image_path, prompt)
            # 清理可能的 markdown 代码块标记
            response = response.strip()
            if response.startswith("```"):
                response = re.sub(r'^```json?\s*\n', '', response)
                response = re.sub(r'\n```\s*$', '', response)

            data = json.loads(response)
            data["page"] = page_num
            return data
        except json.JSONDecodeError as e:
            print(f"⚠️  第 {page_num} 页 JSON 解析失败，降级为快速模式")
            return self._process_page_fast(image_path, page_num)
        except Exception as e:
            print(f"❌ 第 {page_num} 页处理失败: {e}")
            return {
                "page": page_num,
                "elements": [
                    {"type": "paragraph", "content": f"[第 {page_num} 页处理失败: {e}]"}
                ]
            }

    def generate_word(self, pages_data: List[Dict], output_path: str):
        """生成 Word 文档

        Args:
            pages_data: 页面数据列表
            output_path: 输出文件路径
        """
        doc = Document()

        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        print(f"\n📝 生成 Word 文档...")

        for page_data in tqdm(pages_data, desc="写入页面", unit="页"):
            self._add_page_to_doc(doc, page_data)

            # 除了最后一页，其他页添加分页符
            if page_data != pages_data[-1]:
                doc.add_page_break()

        # 保存文档
        doc.save(output_path)
        print(f"✅ Word 文档已保存：{output_path}")

    def _add_page_to_doc(self, doc: Document, page_data: Dict):
        """添加页面内容到文档"""
        layout = page_data.get('layout', 'single_column')

        # 处理分栏布局
        if layout == 'two_column':
            section = doc.sections[-1]
            # 注意：python-docx 对分栏支持有限，这里只是标记
            # 实际效果可能需要手动调整
            pass

        # 添加元素
        elements = page_data.get('elements', [])
        for element in elements:
            try:
                self._add_element(doc, element)
            except Exception as e:
                print(f"⚠️  添加元素失败: {e}")
                # 降级处理：添加纯文本
                doc.add_paragraph(f"[元素处理失败: {element.get('type', 'unknown')}]")

    def _add_element(self, doc: Document, element: Dict):
        """添加单个元素"""
        elem_type = element.get('type')

        if elem_type == 'heading':
            self._add_heading(doc, element)
        elif elem_type == 'paragraph':
            self._add_paragraph(doc, element)
        elif elem_type == 'table':
            self._add_table(doc, element)
        elif elem_type == 'list':
            self._add_list(doc, element)
        else:
            # 未知类型，作为段落处理
            content = element.get('content', '')
            if content:
                doc.add_paragraph(content)

    def _add_heading(self, doc: Document, element: Dict):
        """添加标题"""
        level = element.get('level', 1)
        content = element.get('content', '')

        if not content:
            return

        # 添加标题
        heading = doc.add_heading(content, level=level)

        # 应用样式
        if self.mode == "precise" and 'style' in element:
            self._apply_style(heading, element['style'])

    def _add_paragraph(self, doc: Document, element: Dict):
        """添加段落"""
        content = element.get('content', '')

        if not content:
            return

        # 添加段落
        paragraph = doc.add_paragraph(content)

        # 应用样式
        if self.mode == "precise" and 'style' in element:
            self._apply_style(paragraph, element['style'])

    def _add_table(self, doc: Document, element: Dict):
        """添加表格"""
        rows_data = element.get('rows', [])

        if not rows_data or not rows_data[0]:
            return

        rows = len(rows_data)
        cols = len(rows_data[0])

        # 创建表格
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'

        # 填充数据
        for i, row_data in enumerate(rows_data):
            for j, cell_data in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = str(cell_data)

        # 应用表格样式
        if self.mode == "precise" and 'style' in element:
            style = element['style']
            if style.get('has_header', False) and rows > 0:
                # 加粗表头
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

    def _add_list(self, doc: Document, element: Dict):
        """添加列表"""
        items = element.get('items', [])
        list_type = element.get('list_type', 'bullet')  # bullet 或 number

        style = 'List Bullet' if list_type == 'bullet' else 'List Number'

        for item in items:
            doc.add_paragraph(str(item), style=style)

    def _apply_style(self, paragraph, style: Dict):
        """应用段落样式"""
        if not style:
            return

        # 确保段落有内容
        if not paragraph.runs:
            paragraph.add_run()

        # 应用到所有 run
        for run in paragraph.runs:
            font = run.font

            # 字体
            if 'font' in style:
                font.name = style['font']
                # 设置中文字体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font'])

            # 字号
            if 'size' in style:
                font.size = Pt(style['size'])

            # 粗体
            if 'bold' in style:
                font.bold = style['bold']

            # 斜体
            if 'italic' in style:
                font.italic = style['italic']

            # 颜色
            if 'color' in style:
                color = style['color'].lstrip('#')
                try:
                    font.color.rgb = RGBColor(
                        int(color[0:2], 16),
                        int(color[2:4], 16),
                        int(color[4:6], 16)
                    )
                except (ValueError, IndexError):
                    pass  # 颜色格式错误，跳过

        # 对齐方式
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if 'align' in style and style['align'] in align_map:
            paragraph.alignment = align_map[style['align']]
